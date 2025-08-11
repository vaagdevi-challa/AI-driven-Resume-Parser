import os
import json
import re
import traceback
import fitz  # PyMuPDF
import cohere
from dotenv import load_dotenv
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import time
import tempfile
import pandas as pd
import pythoncom
import win32com.client  # For DOC to DOCX conversion
from db import session, init_db
from models import Resume, WorkExperience

# Call once at the start
init_db()

def save_to_database(extracted, file_name):
    resume = Resume(
        file_name=file_name,
        full_name=extracted.get("full_name"),
        email=extracted.get("email"),
        phone_number=extracted.get("phone_number"),
    )

    for exp in extracted.get("work_experience", []):
        we = WorkExperience(
            company_name=exp.get("company_name", ""),
            customer_name=exp.get("customer_name", ""),
            role=exp.get("role", ""),
            duration=exp.get("duration", ""),
            skills_technologies=", ".join(exp.get("skills_technologies", [])) if exp.get("skills_technologies") else "",
            industry_domain=exp.get("industry_domain", ""),
            location=exp.get("location", "")
        )
        resume.work_experiences.append(we)

    session.add(resume)
    session.commit()


load_dotenv()
API_KEY = os.getenv("COHERE_API_KEY")
if not API_KEY:
    raise RuntimeError("COHERE_API_KEY not set in environment variables or .env file.")
co = cohere.Client(API_KEY, timeout=120)



def convert_doc_to_docx(doc_path):
    pythoncom.CoInitialize()  # Needed if running outside main thread
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    docx_path = os.path.splitext(doc_path)[0] + ".converted.docx"
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, FileFormat=16)  # wdFormatDocumentDefault = 16
        doc.Close()
        return docx_path
    finally:
        word.Quit()

# def extract_text_from_file(filepath):
#     ext = filepath.lower().split(".")[-1]
#     if ext == "pdf":
#         doc = fitz.open(filepath)
#         return "\n".join(page.get_text() for page in doc)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     elif ext == "doc":
#         print(f"🔁 Converting {filepath} to .docx")
#         docx_path = convert_doc_to_docx(filepath)
#         text = extract_text_from_docx(docx_path)
#         os.remove(docx_path)  # Clean up the temporary .docx file
#         return text
#     else:
#         raise ValueError("Unsupported file type: " + ext)


import pdfplumber
import mammoth
from docx import Document

def extract_text_from_pdf(filepath):
    with pdfplumber.open(filepath) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf)

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_doc(filepath):
    with open(filepath, "rb") as f:
        result = mammoth.convert_to_text(f)
        return result.value

def extract_text_from_file(filepath):
    ext = filepath.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    elif ext == "docx":
        return extract_text_from_docx(filepath)
    elif ext == "doc":
        return extract_text_from_doc(filepath)
    else:
        raise ValueError("Unsupported file type: " + ext)

def build_prompt(resume_text: str) -> str:
    return f"""
You are an expert resume data extraction system. Your sole output MUST be a single, valid JSON object. 
Extract the following information from the provided resume text.

The JSON object must have these top-level keys:
1.  "Full Name": The full name of the person. If not found, use JSON null.
2.  "Email": The primary email address. If not found, use JSON null.
3.  "Phone Number": The primary phone number. If not found, use JSON null.
4.  "Work Experience": An array of objects. Each object represents a distinct work experience.

Each object must have:
- "Company Name", "Customer Name", "Role", "Duration", "Skills/Technologies", "Industry/Domain", "Location"

Format must match exactly:
{{
  "Full Name": "Jane Doe",
  "Email": "jane@example.com",
  "Phone Number": "123-456-7890",
  "Work Experience": [
    {{
      "Company Name": "Company A",
      "Customer Name": "Client X",
      "Role": "Developer",
      "Duration": "Jan 2020 - Dec 2021",
      "Skills/Technologies": ["Python", "Django"],
      "Industry/Domain": "IT",
      "Location": "Remote"
    }}
  ]
}}

If resume is invalid or unextractable:
{{
  "Full Name": null,
  "Email": null,
  "Phone Number": null,
  "Work Experience": []
}}

Resume Text:
---
{resume_text}
---
End of Resume Text. Output JSON object:
"""

def clean_json_string(json_str: str) -> str:
    json_str = json_str.strip()
    match = re.search(r"\{.*\}", json_str, re.DOTALL)
    if match:
        json_str = match.group(0)
    json_str = re.sub(r",\s*([}\]])", r"\1", json_str)
    return json_str

def extract_resume_data(resume_text: str):
    if not resume_text.strip():
        return {
            "full_name": None,
            "email": None,
            "phone_number": None,
            "work_experience": []
        }

    prompt = build_prompt(resume_text)
    response = co.generate(model="command-r-plus", prompt=prompt, temperature=0.2)
    raw_output = response.generations[0].text
    json_str = clean_json_string(raw_output)

    try:
        parsed = json.loads(json_str)
    except:
        return extract_resume_data_chunked(resume_text)

    work_exp = parsed.get("Work Experience", [])
    mapped_work_exp = []
    if isinstance(work_exp, list):
        for exp in work_exp:
            if isinstance(exp, dict):
                mapped_work_exp.append({
                    "company_name": exp.get("Company Name", "N/A"),
                    "customer_name": exp.get("Customer Name", "N/A"),
                    "role": exp.get("Role", "N/A"),
                    "duration": exp.get("Duration", "N/A"),
                    "skills_technologies": [s for s in (exp.get("Skills/Technologies") or []) if isinstance(s, str) and s.strip()],

                    "industry_domain": exp.get("Industry/Domain", "N/A"),
                    "location": exp.get("Location", "N/A")
                })

    return {
        "full_name": parsed.get("Full Name"),
        "email": parsed.get("Email"),
        "phone_number": parsed.get("Phone Number"),
        "work_experience": mapped_work_exp
    }

def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    content = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                content.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    content.append(row_text)
    return "\n".join(content)

def extract_text_from_file(filepath):
    ext = filepath.lower().split(".")[-1]
    if ext == "pdf":
        doc = fitz.open(filepath)
        return "\n".join(page.get_text() for page in doc)
    elif ext == "docx":
        return extract_text_from_docx(filepath)
    elif ext == "doc":
        print(f"🔁 Converting {filepath} to .docx")
        docx_path = convert_doc_to_docx(filepath)
        text = extract_text_from_docx(docx_path)
        os.remove(docx_path)
        return text
    else:
        raise ValueError("Unsupported file type: " + ext)

def extract_email_from_lines(lines):
    for line in lines:
        match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', line)
        if match:
            return match.group(0)
    return None

def extract_phone_from_lines(lines):
    for line in lines:
        match = re.search(r'(\+?\d[\d\-\(\) ]{7,}\d)', line)
        if match:
            return match.group(0)
    return None

def parse_resume(file_bytes, filename):
    ext = filename.lower().split(".")[-1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        text = extract_text_from_file(tmp_path)
        result = extract_resume_data(text)

        if not result.get("full_name") and not result.get("email") and not result.get("phone_number"):
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            result["email"] = extract_email_from_lines(lines)
            result["phone_number"] = extract_phone_from_lines(lines)

        return result
    finally:
        os.remove(tmp_path)

def split_experience_sections(resume_text, max_chunks=5):
    sections = re.split(r'\n(?=Experience|Work History|Work Experience|Professional Experience|Employment History|Professional Background)', resume_text, flags=re.IGNORECASE)
    head = sections[0]
    experiences = sections
    chunks = []
    chunk = []
    for exp in experiences:
        chunk.append(exp)
        if len(chunk) >= max_chunks:
            chunks.append('\n'.join(chunk))
            chunk = []
    if chunk:
        chunks.append('\n'.join(chunk))
    return head, chunks

def extract_resume_data_chunked(resume_text: str, chunk_size=5, rate_limit_seconds=5):
    if not resume_text.strip():
        return {
            "full_name": None,
            "email": None,
            "phone_number": None,
            "work_experience": []
        }

    head, exp_chunks = split_experience_sections(resume_text, max_chunks=chunk_size)
    prompt = build_prompt(head)
    response = co.generate(model="command-r-plus", prompt=prompt, temperature=0.2)
    raw_output = response.generations[0].text
    json_str = clean_json_string(raw_output)
    try:
        parsed = json.loads(json_str)
    except Exception:
        return {
            "full_name": None,
            "email": None,
            "phone_number": None,
            "work_experience": []
        }

    full_name = parsed.get("Full Name")
    email = parsed.get("Email")
    phone_number = parsed.get("Phone Number")
    mapped_work_exp = []

    for chunk in exp_chunks:
        time.sleep(rate_limit_seconds)
        prompt = build_prompt(chunk)
        response = co.generate(model="command-r-plus", prompt=prompt, temperature=0.2)
        raw_output = response.generations[0].text
        json_str = clean_json_string(raw_output)
        try:
            parsed_chunk = json.loads(json_str)
        except Exception:
            continue

        work_exp = parsed_chunk.get("Work Experience", [])
        if isinstance(work_exp, list):
            for exp in work_exp:
                if isinstance(exp, dict):
                    mapped_work_exp.append({
                        "company_name": exp.get("Company Name", "N/A"),
                        "customer_name": exp.get("Customer Name", "N/A"),
                        "role": exp.get("Role", "N/A"),
                        "duration": exp.get("Duration", "N/A"),
                        "skills_technologies": [s for s in (exp.get("Skills/Technologies") or []) if isinstance(s, str)],
                        "industry_domain": exp.get("Industry/Domain", "N/A"),
                        "location": exp.get("Location", "N/A")
                    })

    return {
        "full_name": full_name,
        "email": email,
        "phone_number": phone_number,
        "work_experience": mapped_work_exp
    }

def process_resumes_in_folder(folder_path, output_csv="extracted_resumes.csv"):
    supported_extensions = (".pdf", ".docx",".doc")
    data = []
    for filename in os.listdir(folder_path):
        if not filename.lower().endswith(supported_extensions):
            continue

        filepath = os.path.join(folder_path, filename)
        try:
            with open(filepath, "rb") as f:
                file_bytes = f.read()

            print(f"Processing: {filename}")
            extracted = parse_resume(file_bytes, filename)
            if not extracted:
                continue

            entry = {
                "File Name": filename,
                "Full Name": extracted.get("full_name"),
                "Email": extracted.get("email"),
                "Phone Number": extracted.get("phone_number"),
            }

            work_experiences = extracted.get("work_experience", [])
            for i, exp in enumerate(work_experiences, 1):
                entry[f"Company {i}"] = exp.get("company_name", "")
                entry[f"Customer {i}"] = exp.get("customer_name", "")
                entry[f"Role {i}"] = exp.get("role", "")
                entry[f"Duration {i}"] = exp.get("duration", "")
                entry[f"Skills {i}"] = ", ".join(exp.get("skills_technologies", []))
                entry[f"Industry {i}"] = exp.get("industry_domain", "")
                entry[f"Location {i}"] = exp.get("location", "")

                data.append(entry)

        
            save_to_database(extracted, filename)


        except Exception as e:
            print(f"❌ Failed to process {filename}: {e}")
            traceback.print_exc()

    # for filename in os.listdir(folder_path):
    #     if not filename.lower().endswith(supported_extensions):
    #         continue

    #     filepath = os.path.join(folder_path, filename)
    #     try:
    #         with open(filepath, "rb") as f:
    #             file_bytes = f.read()

    #         print(f"Processing: {filename}")
    #         extracted = parse_resume(file_bytes, filename)
    #         if not extracted:
    #             continue

    #         entry = {
    #             "File Name": filename,
    #             "Full Name": extracted.get("full_name"),
    #             "Email": extracted.get("email"),
    #             "Phone Number": extracted.get("phone_number"),
    #         }

    #         work_experiences = extracted.get("work_experience", [])
    #         for i, exp in enumerate(work_experiences, 1):
    #             entry[f"Company {i}"] = exp.get("company_name", "")
    #             entry[f"Customer {i}"] = exp.get("customer_name", "")
    #             entry[f"Role {i}"] = exp.get("role", "")
    #             entry[f"Duration {i}"] = exp.get("duration", "")
    #             entry[f"Skills {i}"] = ", ".join(exp.get("skills_technologies", []))
    #             entry[f"Industry {i}"] = exp.get("industry_domain", "")
    #             entry[f"Location {i}"] = exp.get("location", "")

    #         data.append(entry)

    #     except Exception as e:
    #         print(f"❌ Failed to process {filename}: {e}")
    #         traceback.print_exc()

    if data:
        df = pd.DataFrame(data)
        df.to_csv(output_csv, index=False)
        print(f"✅ Data written to: {output_csv}")
    else:
        print("⚠️ No resume data extracted.")


# Example run (change the folder path accordingly)
if __name__ == "__main__":
    init_db()
    folder_path = input()  # <- Change this to your folder
    process_resumes_in_folder(folder_path, "resume_output.csv")
